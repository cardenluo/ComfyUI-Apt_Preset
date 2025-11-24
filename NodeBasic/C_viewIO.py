
from nodes import MAX_RESOLUTION, SaveImage, common_ksampler
import torch
import os
import sys
import folder_paths
import random
from pathlib import Path
from PIL.PngImagePlugin import PngInfo
from comfy import latent_formats
import json
import latent_preview
from comfy.cli_args import args
import numpy as np
import inspect
import re
import traceback
import itertools
import comfy
from server import PromptServer
from PIL import Image, ImageOps, ImageSequence
import node_helpers
import hashlib
import ast
import io
import base64
from typing import List, Dict, Any,Tuple,Optional
import glob
import torch.nn.functional as F




from ..main_unit import *
from ..office_unit import ImageCompositeMasked



#---------------------ÂÆâÂÖ®ÂØºÂÖ•------
try:
    import cv2
    REMOVER_AVAILABLE = True  # ÂØºÂÖ•ÊàêÂäüÊó∂ËÆæÁΩÆ‰∏∫True
except ImportError:
    cv2 = None
    REMOVER_AVAILABLE = False  # ÂØºÂÖ•Â§±Ë¥•Êó∂ËÆæÁΩÆ‰∏∫False








#‰ºòÂÖà‰ªéÂΩìÂâçÊñá‰ª∂ÊâÄÂú®ÁõÆÂΩï‰∏ãÁöÑ comfy Â≠êÁõÆÂΩï‰∏≠Êü•ÊâæÊ®°Âùó
sys.path.insert(0, os.path.join(os.path.dirname(os.path.realpath(__file__)), "comfy"))  

def updateTextWidget(node, widget, text):
    PromptServer.instance.send_sync("view_Data_text_processed", {"node": node, "widget": widget, "text": text})




#region-----------------------Êî∂Á∫≥-------------------------------------------------------#





class view_LatentAdvanced:
    @classmethod
    def INPUT_TYPES(cls):
        return {"required":
                    {"latent": ("LATENT",),
                    "base_model": (["SD15","SDXL"],),
                    "preview_method": (["auto","taesd","latent2rgb"],),
                    },
            "hidden": {"prompt": "PROMPT",
                        "extra_pnginfo": "EXTRA_PNGINFO",
                        "my_unique_id": "UNIQUE_ID",},
                }

    RETURN_TYPES = ("LATENT",)
    RETURN_NAMES = ("latent",)
    OUTPUT_NODE = True
    FUNCTION = "lpreview"
    CATEGORY = "Apt_Preset/üö´Deprecated/üö´"

    def lpreview(self, latent, base_model, preview_method, prompt=None, extra_pnginfo=None, my_unique_id=None):
        previous_preview_method = args.preview_method
        if preview_method == "taesd":
            temp_previewer = latent_preview.LatentPreviewMethod.TAESD
        elif preview_method == "latent2rgb":
            temp_previewer = latent_preview.LatentPreviewMethod.Latent2RGB
        else:
            temp_previewer = latent_preview.LatentPreviewMethod.Auto

        results = list()

        try:
            args.preview_method=temp_previewer
            preview_format = "PNG"
            load_device=comfy.model_management.vae_offload_device()
            latent_format = {"SD15":latent_formats.SD15,
                            "SDXL":latent_formats.SDXL}[base_model]()

            result=[]
            for i in range(len(latent["samples"])):
                x=latent.copy()
                x["samples"] = latent["samples"][i:i+1].clone()
                x_sample = x["samples"]
                x_sample = x_sample /  {"SD15":6,"SDXL":7.5}[base_model]

                img = latent_preview.get_previewer(load_device, latent_format).decode_latent_to_preview(x_sample)
                full_output_folder, filename, counter, subfolder, filename_prefix = folder_paths.get_save_image_path("",folder_paths.get_temp_directory(), img.height, img.width)
                metadata = None
                if not args.disable_metadata:
                    metadata = PngInfo()
                    if prompt is not None:
                        metadata.add_text("prompt", json.dumps(prompt))
                    if extra_pnginfo is not None:
                        for x in extra_pnginfo:
                            metadata.add_text(x, json.dumps(extra_pnginfo[x]))

                file = "latent_"+"".join(random.choice("0123456789") for x in range(8))+".png"
                img.save(os.path.join(full_output_folder, file), pnginfo=metadata, compress_level=4)
                results.append({"filename": file, "subfolder": subfolder, "type": "temp"})

        finally:
            # Restore global changes
            args.preview_method=previous_preview_method

        return {"result": (latent,), "ui": { "images": results } }




class view_mask(SaveImage):
    
    def __init__(self):
        self.output_dir = folder_paths.get_temp_directory()
        self.type = "temp"
        self.prefix_append = "_temp_" + ''.join(random.choice("abcdefghijklmnopqrstupvxyz") for x in range(5))
        self.compress_level = 4

    @classmethod
    def INPUT_TYPES(s):
        return {
            "required": {"mask": ("MASK",), },  
            "hidden": {"prompt": "PROMPT", "extra_pnginfo": "EXTRA_PNGINFO"},
        }

    FUNCTION = "execute"
    CATEGORY = "Apt_Preset/View_IO"
    DESCRIPTION = "show mask"
    
    def execute(self, mask, filename_prefix="ComfyUI", prompt=None, extra_pnginfo=None):
        # Â§ÑÁêÜÂàóË°®Á±ªÂûãÁöÑÈÅÆÁΩ©
        if isinstance(mask, list):
            # Â≠òÂÇ®ÊâÄÊúâÂ§ÑÁêÜÂêéÁöÑÈÅÆÁΩ©
            processed_masks = []
            for m in mask:
                # Á°Æ‰øùÊØè‰∏™ÂÖÉÁ¥†ÈÉΩÊòØÂº†Èáè
                if isinstance(m, torch.Tensor):
                    processed = self.process_single_mask(m)
                    processed_masks.append(processed)
            
            # ÂêàÂπ∂ÊâÄÊúâÈÅÆÁΩ©‰∏∫‰∏Ä‰∏™ÊâπÊ¨°
            if processed_masks:
                preview = torch.cat(processed_masks, dim=0)
            else:
                # Â§ÑÁêÜÁ©∫ÂàóË°®ÊÉÖÂÜµ
                return {"ui": {"images": []}}
        # Â§ÑÁêÜÂçï‰∏™Âº†ÈáèÈÅÆÁΩ©
        elif isinstance(mask, torch.Tensor):
            preview = self.process_single_mask(mask)
        else:
            # Â§ÑÁêÜÂÖ∂‰ªñ‰∏çÊîØÊåÅÁöÑÁ±ªÂûã
            return {"ui": {"images": []}}
        
        return self.save_images(preview, filename_prefix, prompt, extra_pnginfo)
    
    def process_single_mask(self, mask_tensor):
        """Â§ÑÁêÜÂçï‰∏™ÈÅÆÁΩ©Âº†ÈáèÔºåËΩ¨Êç¢‰∏∫Ê≠£Á°ÆÁöÑÈ¢ÑËßàÊ†ºÂºè"""
        # Ê†πÊçÆÂº†ÈáèÁª¥Â∫¶ËøõË°å‰∏çÂêåÂ§ÑÁêÜ
        if mask_tensor.dim() == 2:  # ÂΩ¢Áä∂‰∏∫ (H, W)
            # Ê∑ªÂä†ÊâπÊ¨°ÂíåÈÄöÈÅìÁª¥Â∫¶: (1, 1, H, W) -> ËΩ¨Êç¢Âêé (1, H, W, 3)
            return mask_tensor.unsqueeze(0).unsqueeze(0).movedim(1, -1).expand(-1, -1, -1, 3)
        elif mask_tensor.dim() == 3:  # ÂΩ¢Áä∂‰∏∫ (B, H, W) Êàñ (1, H, W)
            # Ê∑ªÂä†ÈÄöÈÅìÁª¥Â∫¶Âπ∂ËΩ¨Êç¢: (B, 1, H, W) -> (B, H, W, 3)
            return mask_tensor.unsqueeze(1).movedim(1, -1).expand(-1, -1, -1, 3)
        else:  # ÂÖ∂‰ªñÁª¥Â∫¶Ôºå‰ΩøÁî®reshapeÁ°Æ‰øùÂÖºÂÆπÊÄß
            return mask_tensor.reshape((-1, 1, mask_tensor.shape[-2], mask_tensor.shape[-1])).movedim(1, -1).expand(-1, -1, -1, 3)
    


class view_combo:     # web_node/view_Data_text.js

    @classmethod
    def INPUT_TYPES(s):
        return {"required": {
                    "prompt": ("STRING", {"multiline": True, "default": "text"}),
                    "start_index": ("INT", {"default": 0, "min": 0, "max": 9999}),
                    "max_rows": ("INT", {"default": 1000, "min": 1, "max": 9999}),
                    },
            "hidden":{
                "workflow_prompt": "PROMPT", "my_unique_id": "UNIQUE_ID"
            }
        }

    RETURN_TYPES = (any_type, any_type)
    RETURN_NAMES = ("STRING", "COMBO")
    OUTPUT_IS_LIST = (True, True)
    FUNCTION = "generate_strings"
    CATEGORY = "Apt_Preset/View_IO"

    def generate_strings(self, prompt, start_index, max_rows, workflow_prompt=None, my_unique_id=None):
        lines = prompt.split('\n')

        start_index = max(0, min(start_index, len(lines) - 1))
        end_index = min(start_index + max_rows, len(lines))
        rows = lines[start_index:end_index]

        return (rows, rows)




class view_node_Script:
    def __init__(self):
        self.node_list = []
        self.custom_node_list = []
        self.update_node_list()

    def update_node_list(self):
        try:
            import nodes
            self.node_list = []
            self.custom_node_list = []
            
            for node_name, node_class in nodes.NODE_CLASS_MAPPINGS.items():
                try:
                    module = inspect.getmodule(node_class)
                    module_path = getattr(module, '__file__', '')
                    is_custom = 'custom_nodes' in module_path

                    node_info = {
                        'name': node_name,
                        'class_name': node_class.__name__,
                        'category': getattr(node_class, 'CATEGORY', 'Uncategorized'),
                        'description': getattr(node_class, 'DESCRIPTION', ''),
                        'is_custom': is_custom
                    }
                    
                    self.node_list.append(node_info)
                    if is_custom:
                        self.custom_node_list.append(node_info)
                except Exception as e:
                    logging.error(f"Error processing node {node_name}: {str(e)}")
                    continue
            
            self.node_list.sort(key=lambda x: x['name'])
            self.custom_node_list.sort(key=lambda x: x['name'])
            
        except Exception as e:
            logging.error(f"Error updating node list: {str(e)}")
            traceback.print_exc()

    @classmethod
    def INPUT_TYPES(cls):
        try:
            import nodes
            node_names = sorted(list(nodes.NODE_CLASS_MAPPINGS.keys()))
            if not node_names:
                node_names = ["No nodes found"]
                
            return {
                "required": {
                    "selected_node": (node_names, {
                        "default": node_names[0]
                    }),
                    "search": ("STRING", {
                        "default": "",
                        "multiline": False
                    }),
                    "show_all": ("BOOLEAN", {
                        "default": True,
                        "label": "Show All Nodes"
                    }),
                    "refresh_list": ("BOOLEAN", {
                        "default": False,
                        "label": "Refresh Node List"
                    })
                }
            }
        except Exception as e:
            print(f"Error in INPUT_TYPES: {str(e)}")
            return {
                "required": {
                    "search": ("STRING", {"default": "", "multiline": False}),
                    "show_all": ("BOOLEAN", {"default": True, "label": "Show All Nodes"}),
                    "refresh_list": ("BOOLEAN", {"default": False, "label": "Refresh Node List"})
                }
            }

    RETURN_TYPES = ("STRING",)
    RETURN_NAMES = ("node_source",)
    FUNCTION = "find_script"
    CATEGORY = "Apt_Preset/View_IO"

    def get_node_source_code(self, node_name):
        try:
            import nodes
            import inspect
            import os

            node_class = nodes.NODE_CLASS_MAPPINGS.get(node_name)
            if not node_class:
                return f"Node '{node_name}' not found"

            module = inspect.getmodule(node_class)
            if not module:
                return f"Could not find module for {node_name}"

            try:
                file_path = inspect.getfile(module)
            except TypeError:
                return f"Could not determine file path for {node_name}"

            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    file_content = f.read()
            except Exception as e:
                return f"Error reading file: {str(e)}"

            class_def = f"class {node_class.__name__}:"
            class_start = file_content.find(class_def)
            
            if class_start == -1:
                return f"Could not find class definition for {node_name}"

            lines = file_content[class_start:].split('\n')
            class_lines = []
            indent_level = None

            for line in lines:
                if indent_level is None:
                    if line.strip().startswith('class'):
                        indent_level = len(line) - len(line.lstrip())
                    continue

                current_indent = len(line) - len(line.lstrip())
                if current_indent <= indent_level and line.strip():
                    break

                class_lines.append(line)

            source_output = f"=== Node: {node_name} ===\n"
            source_output += f"File: {file_path}\n\n"
            source_output += "=== Source Code ===\n"
            source_output += "\n".join(class_lines)

            return source_output

        except Exception as e:
            return f"Error retrieving source code: {str(e)}"

    def find_script(self, selected_node, search, show_all, refresh_list):
        try:
            if refresh_list:
                self.update_node_list()

            if selected_node:
                source_code = self.get_node_source_code(selected_node)
                return (source_code,)
            return ("Please select a node to view its source code",)

        except Exception as e:
            logging.error(f"Error in find_script: {str(e)}")
            traceback.print_exc()
            return (traceback.format_exc(),)









class IPA_clip_vision:
    @classmethod
    def INPUT_TYPES(s):
        return {"required": { 
            "clip_name": (folder_paths.get_filename_list("clip_vision"), ),
            "image": ("IMAGE",),
        }}
    RETURN_TYPES = ("CLIP_VISION_OUTPUT",)
    FUNCTION = "combined_process"

    CATEGORY = "Apt_Preset/chx_tool/chx_IPA"

    def combined_process(self, clip_name, image):
        # Âä†ËΩΩ CLIP Vision Ê®°Âûã
        clip_path = folder_paths.get_full_path_or_raise("clip_vision", clip_name)
        clip_vision = comfy.clip_vision.load(clip_path)
        if clip_vision is None:
            raise RuntimeError("ERROR: clip vision file is invalid and does not contain a valid vision model.")
        
        output = clip_vision.encode_image(image, crop="center")
        return (output,)



#region----------------------cache -clear all-------------------------------------------------------#

class TaggedCache:
    def __init__(self, tag_settings: Optional[dict]=None):
        self._tag_settings = tag_settings or {}  
        self._data = {}

    def __getitem__(self, key):
        for tag_data in self._data.values():
            if key in tag_data:
                return tag_data[key]
        raise KeyError(f'Key `{key}` does not exist')

    def __setitem__(self, key, value: tuple):

        for tag_data in self._data.values():
            if key in tag_data:
                tag_data.pop(key, None)
                break

        tag = value[0]
        if tag not in self._data:

            try:
                from cachetools import LRUCache
                default_size = 20
                if 'ckpt' in tag:
                    default_size = 5
                elif tag in ['latent', 'image']:
                    default_size = 100
                self._data[tag] = LRUCache(maxsize=self._tag_settings.get(tag, default_size))
            except (ImportError, ModuleNotFoundError):
                self._data[tag] = {}
        self._data[tag][key] = value

    def __delitem__(self, key):
        for tag_data in self._data.values():
            if key in tag_data:
                del tag_data[key]
                return
        raise KeyError(f'Key `{key}` does not exist')

    def __contains__(self, key):
        return any(key in tag_data for tag_data in self._data.values())

    def items(self):
        yield from itertools.chain(*map(lambda x :x.items(), self._data.values()))

    def get(self, key, default=None):
        for tag_data in self._data.values():
            if key in tag_data:
                return tag_data[key]
        return default

    def clear(self):
        self._data = {}

cache_settings = {}
cache = TaggedCache(cache_settings)
cache_count = {}

def update_cache(k, tag, v):
    cache[k] = (tag, v)
    cnt = cache_count.get(k)
    if cnt is None:
        cnt = 0
        cache_count[k] = cnt
    else:
        cache_count[k] += 1
def remove_cache(key):
    global cache
    if key == '*':
        cache = TaggedCache(cache_settings)
    elif key in cache:
        del cache[key]
    else:
        print(f"invalid {key}")



class IO_clear_cache:
    @classmethod
    def INPUT_TYPES(s):
        return {"required": {
            "anything": (any_type, {}),
        }, "optional": {},
            "hidden": {"unique_id": "UNIQUE_ID", "extra_pnginfo": "EXTRA_PNGINFO", }
        }

    RETURN_TYPES = (any_type,)
    RETURN_NAMES = ("output",)
    OUTPUT_NODE = True
    FUNCTION = "empty_cache"
    CATEGORY = "Apt_Preset/View_IO/üò∫backup"

    def empty_cache(self, anything, unique_id=None, extra_pnginfo=None):
        remove_cache('*')
        return (anything,)

#endregion-----------------------clear all-------------------------------------------------------#



class IO_inputbasic:
    def __init__(self):
        pass
    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {
                "input": ("STRING", {
                    "multiline": True,
                    "default": ""
                }),
            }
        }
    RETURN_TYPES = ("INT", "FLOAT", "STRING")
    RETURN_NAMES = ("int", "float", "string")
    FUNCTION = "convert_number_types"
    CATEGORY = "Apt_Preset/View_IO/üò∫backup"
    def convert_number_types(self, input):
        try:
            float_num = float(input)
            int_num = int(float_num)
            str_num = input
        except ValueError:
            return (None, None, input)
        return (int_num, float_num, str_num)


class view_Data:   

    @classmethod
    def INPUT_TYPES(s):
        return {
            "optional": {
                "any": (anyType, {"forceInput": True}),
                "data": ("STRING", {"default": "", "multiline": True}),
            },
            "hidden": {
                "unique_id": "UNIQUE_ID",
            },
        }

    RETURN_TYPES = (anyType,)  # Âè™ÈúÄË¶Å‰∏Ä‰∏™ËæìÂá∫Á´ØÂè£
    RETURN_NAMES = ("-",)  # ËæìÂá∫ÂêçÁß∞
    INPUT_IS_LIST = (True,)
    OUTPUT_NODE = True

    CATEGORY = "Apt_Preset/View_IO"
    FUNCTION = "process"

    def process(self, data, unique_id, any=None):
        displayText = self.render(any)

        updateTextWidget(unique_id, "data", displayText)
        if isinstance(any, list) and len(any) == 1:
            return {"ui": {"data": displayText}, "result": (any[0],)}
        else:
            return {"ui": {"data": displayText}, "result": (any,)}

    def render(self, any):
        if not isinstance(any, list):
            return str(any)

        listLen = len(any)

        if listLen == 0:
            return ""

        if listLen == 1:
            return str(any[0])

        result = "List:\n"

        for i, element in enumerate(any):
            result += f"- {str(any[i])}\n"

        return result




class view_GetLength:
    def __init__(self):
        pass

    @classmethod
    def INPUT_TYPES(s):
        return {
            "required": {
                "ANY" : (ANY_TYPE, {}), 
            },
        }
    
    TITLE = "Get Length"
    RETURN_TYPES = ("INT", )
    RETURN_NAMES = ("length", )
    FUNCTION = "run"
    CATEGORY = "Apt_Preset/View_IO"
    OUTPUT_NODE = True

    def run(self, ANY):
        length = len(ANY)
        return { "ui": {"text": (f"{length}",)}, "result": (length, ) }




class view_GetShape:
    def __init__(self):
        pass

    @classmethod
    def INPUT_TYPES(s):
        return {
            "required": {
                "tensor" : ("IMAGE,LATENT,MASK", {}), 
            },
            "hidden": {
                "unique_id": "UNIQUE_ID",
                "prompt": "PROMPT", 
                "extra_pnginfo": "EXTRA_PNGINFO",
            },
        }
    
    TITLE = "Get Shape"
    RETURN_TYPES = ("INT", "INT", "INT", "INT")
    RETURN_NAMES = ("width", "height", "batch_size", "channels")
    FUNCTION = "run"
    CATEGORY = "Apt_Preset/View_IO"
    OUTPUT_NODE = True

    @classmethod
    def VALIDATE_INPUTS(cls, input_types):
        return True

    def run(self, tensor, unique_id, prompt, extra_pnginfo):  
        node_list = extra_pnginfo["workflow"]["nodes"]
        cur_node = next(n for n in node_list if str(n["id"]) == unique_id)
        link_id = cur_node["inputs"][0]["link"]
        link = next(l for l in extra_pnginfo["workflow"]["links"] if l[0] == link_id)
        in_node_id, in_socket_id = link[1], link[2]
        in_node = next(n for n in node_list if n["id"] == in_node_id)
        input_type = in_node["outputs"][in_socket_id]["type"]
        
        B, H, W, C = 1, 1, 1, 1
        
        if input_type == "IMAGE":
            B, H, W, C = tensor.shape
        elif input_type == "LATENT" or (type(tensor) is dict and "samples" in tensor):
            samples_shape = tensor["samples"].shape
            if len(samples_shape) == 4:
                B, C, H, W = samples_shape
            elif len(samples_shape) == 3:
                B, H, W = samples_shape
                C = 4
            else:
                print(f"Unexpected latent shape: {samples_shape}")
                B, C, H, W = 1, 4, 64, 64
            
            H *= 8
            W *= 8
        else:
            shape = tensor.shape
            if len(shape) == 2:
                H, W = shape
            elif len(shape) == 3:
                B, H, W = shape
            elif len(shape) == 4:
                if shape[3] <= 4:
                    B, H, W, C = tensor.shape
                else:
                    B, C, H, W = shape
        
        return { "ui": {"text": (f"{W}, {H}, {B}, {C}",)}, "result": (W, H, B, C) }



class view_GetWidgetsValues:
    def __init__(self):
        pass

    @classmethod
    def INPUT_TYPES(s):
        return {
            "required": {
                "ANY" : (ANY_TYPE, {}), 
            },
            "hidden": {
                "unique_id": "UNIQUE_ID",
                "prompt": "PROMPT", 
                "extra_pnginfo": "EXTRA_PNGINFO",
            },
        }
    
    TITLE = "Get Widgets Values"
    RETURN_TYPES = ("LIST", )
    RETURN_NAMES = ("LIST", )
    OUTPUT_NODE = True
    FUNCTION = "run"
    CATEGORY = "Apt_Preset/View_IO"
    OUTPUT_NODE = True

    def run(self, ANY, unique_id, prompt, extra_pnginfo):
        node_list = extra_pnginfo["workflow"]["nodes"]  # list of dict including id, type
        cur_node = next(n for n in node_list if str(n["id"]) == unique_id)
        link_id = cur_node["inputs"][0]["link"]
        link = next(l for l in extra_pnginfo["workflow"]["links"] if l[0] == link_id)
        in_node_id, in_socket_id = link[1], link[2]
        in_node = next(n for n in node_list if n["id"] == in_node_id)
        return { "ui": {"text": (f"{in_node['widgets_values']}",)}, "result": (in_node["widgets_values"], ) }




#endregion-----------------------Êóß-------------------------------------------------------#.


class view_Mask_And_Img(SaveImage):
    def __init__(self):
        self.output_dir = folder_paths.get_temp_directory()
        self.type = "temp"
        self.prefix_append = "_temp_" + ''.join(random.choice("abcdefghijklmnopqrstupvxyz") for x in range(5))
        self.compress_level = 4

    @classmethod
    def INPUT_TYPES(s):
        return {
            "required": {
                "mask_opacity": ("FLOAT", {"default": 1.0, "min": 0.0, "max": 1.0, "step": 0.01}),
            },
            "optional": {
                "image": ("IMAGE",),
                "mask": ("MASK",),                
            },
            "hidden": {"prompt": "PROMPT", "extra_pnginfo": "EXTRA_PNGINFO"},
        }


    FUNCTION = "execute"
    CATEGORY = "Apt_Preset/View_IO"


    def execute(self, mask_opacity, filename_prefix="ComfyUI", image=None, mask=None, prompt=None, extra_pnginfo=None):
        if mask is not None and image is None:
            preview = mask.reshape((-1, 1, mask.shape[-2], mask.shape[-1])).movedim(1, -1).expand(-1, -1, -1, 3)
        elif mask is None and image is not None:
            if image.shape[-1] == 4:
                image = image[..., :3]
            preview = image
        elif mask is not None and image is not None:
            if image.shape[-1] == 4:
                image = image[..., :3]
            mask_adjusted = mask * mask_opacity
            mask_image = mask.reshape((-1, 1, mask.shape[-2], mask.shape[-1])).movedim(1, -1).expand(-1, -1, -1, 3).clone()
            color_list = [0, 0, 0]
            mask_image[:, :, :, 0] = color_list[0] / 255
            mask_image[:, :, :, 1] = color_list[1] / 255
            mask_image[:, :, :, 2] = color_list[2] / 255
            preview, = ImageCompositeMasked.composite(self, image, mask_image, 0, 0, True, mask_adjusted)
        else:
            preview = torch.zeros((1, 64, 64, 3), dtype=torch.float32, device="cpu")

        return self.save_images(preview, filename_prefix, prompt, extra_pnginfo)





class IO_adjust_image:
    @classmethod
    def INPUT_TYPES(s):
        input_dir = folder_paths.get_input_directory()
        files = [f.name for f in Path(input_dir).iterdir() if f.is_file()]
        return {
            "required": {
                "image": (sorted(files), {"image_upload": True}),
                "max_dimension": ("INT", {"default": 0, "min": 0, "max": 4096, "step": 8}),
                "size_option": (["No Change", "Custom", "Million Pixels", "Small", "Medium", "Large", 
                                "480P-H(vid 4:3)", "480P-V(vid 3:4)", "720P-H(vid 16:9)", "720P-V(vid 9:16)", "832√ó480", "480√ó832"], 
                                {"default": "No Change"})
            }
        }

    RETURN_TYPES = ("IMAGE", "MASK", "STRING")
    RETURN_NAMES = ("image", "mask", "info")
    FUNCTION = "load_image"
    CATEGORY = "Apt_Preset/üö´Deprecated/üö´"

    def IS_CHANGED(): return float("NaN")

    def load_image(self, image, max_dimension, size_option):
        image_path = folder_paths.get_annotated_filepath(image)
        img = Image.open(image_path)
        W, H = img.size
        aspect_ratio = W / H

        def get_target_size():
            if size_option == "No Change":
                # No resizing or cropping, just return the original size
                return W, H
            elif size_option == "Million Pixels":
                return self._resize_to_million_pixels(W, H)
            elif size_option == "Custom":
                ratio = min(max_dimension / W, max_dimension / H)
                return round(W * ratio), round(H * ratio)
            
            size_options = {
                "Small": (
                    (768, 512) if aspect_ratio >= 1.23 else
                    (512, 768) if aspect_ratio <= 0.82 else
                    (768, 768)
                ),
                "Medium": (
                    (1216, 832) if aspect_ratio >= 1.23 else
                    (832, 1216) if aspect_ratio <= 0.82 else
                    (1216, 1216)
                ),
                "Large": (
                    (1600, 1120) if aspect_ratio >= 1.23 else
                    (1120, 1600) if aspect_ratio <= 0.82 else
                    (1600, 1600)
                ),
                "Million Pixels": self._resize_to_million_pixels(W, H),  # Million Pixels option
                "480P-H(vid 4:3)": (640, 480),  # 480P-H, 640x480
                "480P-V(vid 3:4)": (480, 640),  # 480P-V, 480x640
                "720P-H(vid 16:9)": (1280, 720),  # 720P-H, 1280x720
                "720P-V(vid 9:16)": (720, 1280),  # 720P-V, 720x1280
                "832√ó480": (832, 480),  # 832x480
                "480√ó832": (480, 832),  # 480x832
            }
            return size_options[size_option]
        
        target_width, target_height = get_target_size()
        output_images = []
        output_masks = []

        for frame in ImageSequence.Iterator(img):
            frame = ImageOps.exif_transpose(frame)
            if frame.mode == 'P':
                frame = frame.convert("RGBA")
            elif 'A' in frame.getbands():
                frame = frame.convert("RGBA")
            
            if size_option == "No Change":
                # No resizing, just use the original frame
                image_frame = frame.convert("RGB")
            else:
                if size_option == "Custom" or size_option == "Million Pixels":
                    ratio = min(target_width / W, target_height / H)
                    adjusted_width = round(W * ratio)
                    adjusted_height = round(H * ratio)
                    image_frame = frame.convert("RGB").resize((adjusted_width, adjusted_height), Image.Resampling.BILINEAR)
                else:
                    image_frame = frame.convert("RGB")
                    image_frame = ImageOps.fit(image_frame, (target_width, target_height), method=Image.Resampling.BILINEAR, centering=(0.5, 0.5))

            image_array = np.array(image_frame).astype(np.float32) / 255.0
            output_images.append(torch.from_numpy(image_array)[None,])

            # Process the mask if available
            if 'A' in frame.getbands():
                mask_frame = frame.getchannel('A')
                if size_option == "Custom" or size_option == "Million Pixels":
                    mask_frame = mask_frame.resize((adjusted_width, adjusted_height), Image.Resampling.BILINEAR)
                else:
                    mask_frame = ImageOps.fit(mask_frame, (target_width, target_height), method=Image.Resampling.BILINEAR, centering=(0.5, 0.5))
                mask = np.array(mask_frame).astype(np.float32) / 255.0
                mask = 1. - mask
            else:
                if size_option == "Custom" or size_option == "Million Pixels":
                    mask = np.zeros((adjusted_height, adjusted_width), dtype=np.float32)
                else:
                    mask = np.zeros((target_height, target_width), dtype=np.float32)
            output_masks.append(torch.from_numpy(mask).unsqueeze(0))
        
        output_image = torch.cat(output_images, dim=0) if len(output_images) > 1 else output_images[0]
        output_mask = torch.cat(output_masks, dim=0) if len(output_masks) > 1 else output_masks[0]
        info = f"Image Path: {image_path}\nOriginal Size: {W}x{H}\nAdjusted Size: {target_width}x{target_height}"
        return (output_image, output_mask, info)

    @classmethod
    def VALIDATE_INPUTS(s, image):
        if not folder_paths.exists_annotated_filepath(image):
            return f"Invalid image file: {image}"
        return True
    def _resize_to_million_pixels(self, W, H):
        aspect_ratio = W / H
        target_area = 1000000  # 1 million pixels
        if aspect_ratio > 1:  # Landscape
            width = int(np.sqrt(target_area * aspect_ratio))
            height = int(target_area / width)
        else:  # Portrait
            height = int(np.sqrt(target_area / aspect_ratio))
            width = int(target_area / height)
        width = (width + 7) // 8 * 8
        height = (height + 7) // 8 * 8
        return width, height




class IO_save_image:
    def __init__(self):
        self.type = "output"

    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {
                "image": ("IMAGE", ),
                "file_format": (["png", "webp", "jpg", "tif"],),
                "output_path": ("STRING", {"default": "./output/Apt", "multiline": False}),
                "filename_mid": ("STRING", {"default": "Apt"}),
            },
            "optional": {
                "number_prefix": ("BOOLEAN", {"default": False, "label_on": "ÂâçÁΩÆÁºñÂè∑", "label_off": "ÂêéÁΩÆÁºñÂè∑"}),
                "number_digits": ("INT", {"default": 5, "min": 1, "max": 10, "step": 1, "tooltip": "ÁºñÂè∑‰ΩçÊï∞ÔºåÂ¶ÇËÆæÁΩÆ‰∏∫3Âàô‰∏∫001Ê†ºÂºè"}),
                "save_workflow_as_json": ("BOOLEAN", {"default": False}),
            },
            "hidden": {
                "prompt": "PROMPT",
                "extra_pnginfo": "EXTRA_PNGINFO"
            }
        }

    RETURN_TYPES = ("STRING", )
    RETURN_NAMES = ("out_path",)
    FUNCTION = "save_image"
    OUTPUT_NODE = True
    OUTPUT_IS_LIST = (True,)
    CATEGORY = "Apt_Preset/View_IO"

    @staticmethod
    def find_highest_numeric_value(directory, filename_mid):
        highest_value = -1
        if not os.path.exists(directory):
            return highest_value
        for filename in os.listdir(directory):
            if filename.startswith(filename_mid):
                try:
                    numeric_part = filename[len(filename_mid):]
                    numeric_str = re.search(r'\d+', numeric_part).group()
                    numeric_value = int(numeric_str)
                    if numeric_value > highest_value:
                        highest_value = numeric_value
                except (ValueError, AttributeError):
                    continue
        return highest_value

    def save_image(self, image, file_format, filename_mid="Apt", output_path="", number_prefix=False, number_digits=5,
                   save_workflow_as_json=False, prompt=None, extra_pnginfo=None):
        batch_size = image.shape[0]
        images_list = [image[i:i + 1, ...] for i in range(batch_size)]
        output_dir = folder_paths.get_output_directory()
        output_paths = []

        if isinstance(output_path, str):
            os.makedirs(output_path, exist_ok=True)
            output_paths = [output_path] * batch_size
        elif isinstance(output_path, list) and len(output_path) == batch_size:
            for path in output_path:
                os.makedirs(path, exist_ok=True)
            output_paths = output_path
        else:
            print("Invalid output_path format. Using default output directory.")
            output_paths = [output_dir] * batch_size

        base_dir = output_paths[0]
        counter = self.find_highest_numeric_value(base_dir, filename_mid) + 1
        absolute_paths = []

        for idx, img_tensor in enumerate(images_list):
            output_image = img_tensor.cpu().numpy()
            img_np = np.clip(output_image * 255.0, 0, 255).astype(np.uint8)
            img = Image.fromarray(img_np[0])
            out_path = output_paths[idx]

            numbering = f"{counter + idx:0{number_digits}d}"
            if number_prefix:
                output_filename = f"{numbering}_{filename_mid}"
            else:
                output_filename = f"{filename_mid}_{numbering}"

            resolved_image_path = os.path.join(out_path, f"{output_filename}.{file_format}")
            img_params = {
                'png': {'compress_level': 4},
                'webp': {'method': 6, 'lossless': False, 'quality': 80},
                'jpg': {'quality': 95, 'format': 'JPEG'},
                'tif': {'format': 'TIFF'}
            }
            img.save(resolved_image_path, **img_params[file_format])
            absolute_paths.append(os.path.abspath(resolved_image_path))

            if save_workflow_as_json:
                try:
                    workflow = (extra_pnginfo or {}).get('workflow')
                    if workflow is not None:
                        json_file_path = os.path.join(out_path, f"{output_filename}.json")
                        with open(json_file_path, 'w') as f:
                            json.dump(workflow, f)
                except Exception as e:
                    print(f"Failed to save workflow JSON: {e}")

        return (absolute_paths, )








class IO_input_any:
    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {
                "text": ("STRING", {"multiline": True, "default": "", }),
            },
            "optional": {                
                "delimiter": ("STRING", {
                    "default": "",
                    "multiline": False,
                    "tooltip": "ÂàÜÈöîË°å\\n, ÂàÜÈöîÂà∂Ë°®Á¨¶\\t, ÂàÜÈöîÁ©∫Ê†º\\s"
                }),
                "output_type": (["float", "int", "string", "anytype", "dictionary", "set", "tuple", "boolean"], {"default": "anytype"}),
            }
        }

    RETURN_TYPES = (ANY_TYPE, "LIST")
    RETURN_NAMES = ("data", "list")
    FUNCTION = "process_text"
    CATEGORY = "Apt_Preset/View_IO"
    OUTPUT_IS_LIST = (True, False)

    def process_text(self, text, delimiter="", output_type="anytype"):
        # Â§ÑÁêÜÁâπÊÆäÂàÜÈöîÁ¨¶Ë°®Á§∫
        if delimiter == "\\n":
            delimiter = "\n"
        elif delimiter == "\\t":
            delimiter = "\t"
        elif delimiter == "\\s":
            delimiter = " "
        
        # ÂéªÈô§È¶ñÂ∞æÁ©∫ÁôΩ
        text = text.strip()
        
        # Ê£ÄÊü•ÊòØÂê¶‰∏∫Â∏ÉÂ∞îÂÄºÊ†ºÂºè
        if output_type == "boolean":
            try:
                # Â§ÑÁêÜÂ∏∏ËßÅÁöÑÂ∏ÉÂ∞îÂÄºË°®Á§∫
                if text.lower() in ["true", "yes", "1", "on"]:
                    return ([True], [text])
                elif text.lower() in ["false", "no", "0", "off"]:
                    return ([False], [text])
                else:
                    # ÈùûÊ†áÂáÜÂ∏ÉÂ∞îÂÄºË°®Á§∫ÔºåÂ∞ùËØïËá™Âä®ËΩ¨Êç¢
                    return ([bool(ast.literal_eval(text))], [text])
            except Exception as e:
                print(f"Ëß£ÊûêÂ∏ÉÂ∞îÂÄºÂ§±Ë¥•: {e}")
                return ([False], [text])  # ÈªòËÆ§ËøîÂõû False
        
        # Ê£ÄÊü•ÊòØÂê¶‰∏∫ÁâπÊÆäÁ±ªÂûãÊ†ºÂºèÂπ∂‰ºòÂÖàÂ§ÑÁêÜ
        if output_type == "dictionary" or (output_type == "anytype" and text.startswith("{") and text.endswith("}")):
            try:
                parsed = ast.literal_eval(text)
                if not isinstance(parsed, dict):
                    raise ValueError("Ëß£ÊûêÁªìÊûú‰∏çÊòØÂ≠óÂÖ∏")
                return ([parsed], [text])
            except Exception as e:
                print(f"Ëß£ÊûêÂ≠óÂÖ∏Â§±Ë¥•: {e}")
        
        elif output_type == "set" or (output_type == "anytype" and text.startswith("{") and text.endswith("}") and ":" not in text):
            try:
                # Â§ÑÁêÜÈõÜÂêàÊ†ºÂºèÔºàÈúÄË¶ÅÊ∑ªÂä†Â§ñÂõ¥Êã¨Âè∑‰ª•Á¨¶Âêà Python ËØ≠Ê≥ïÔºâ
                if text == "{}":  # Á©∫ÈõÜÂêà
                    parsed = set()
                else:
                    # ÁßªÈô§È¶ñÂ∞æÊã¨Âè∑Âπ∂Ê∑ªÂä†Â§ñÂõ¥ÂÖÉÁªÑÊã¨Âè∑
                    set_content = text[1:-1]
                    parsed = set(ast.literal_eval(f"({set_content},)"))
                return ([parsed], [text])
            except Exception as e:
                print(f"Ëß£ÊûêÈõÜÂêàÂ§±Ë¥•: {e}")
        
        elif output_type == "tuple" or (output_type == "anytype" and (text.startswith("(") and text.endswith(")") or "," in text)):
            try:
                # Â§ÑÁêÜÂÖÉÁªÑÊ†ºÂºè
                if text == "()":  # Á©∫ÂÖÉÁªÑ
                    parsed = ()
                elif text.endswith(",") and not text.startswith("("):
                    # ÂçïÂÖÉÁ¥†ÂÖÉÁªÑÁâπÊÆäÊ†ºÂºè: "1,"
                    parsed = ast.literal_eval(f"({text})")
                else:
                    parsed = ast.literal_eval(text)
                if not isinstance(parsed, tuple):
                    parsed = (parsed,)  # Á°Æ‰øùÊòØÂÖÉÁªÑ
                return ([parsed], [text])
            except Exception as e:
                print(f"Ëß£ÊûêÂÖÉÁªÑÂ§±Ë¥•: {e}")
        
        # ‰ΩøÁî®ÂàÜÈöîÁ¨¶ÂàÜÂâ≤ÊñáÊú¨
        if delimiter:
            items = text.split(delimiter)
        else:
            # Â¶ÇÊûúÊ≤°ÊúâÊåáÂÆöÂàÜÈöîÁ¨¶Ôºå‰ΩøÁî®ÁÅµÊ¥ªÁöÑÂàÜÈöîÁ¨¶ÂåπÈÖç
            items = re.split(r'[\s,]+', text)
        
        # ÂéªÈô§Á©∫Â≠óÁ¨¶‰∏≤
        items = [item.strip() for item in items if item.strip()]
        
        # ÁîüÊàêÁ±ªÂûãËΩ¨Êç¢ÂêéÁöÑÁªìÊûú
        converted_result = []
        for item in items:
            if output_type == "int":
                try:
                    converted_result.append(int(item))
                except ValueError:
                    converted_result.append(0)  # ËΩ¨Êç¢Â§±Ë¥•Êó∂ÈªòËÆ§‰∏∫0
            elif output_type == "float":
                try:
                    converted_result.append(float(item))
                except ValueError:
                    converted_result.append(0.0)  # ËΩ¨Êç¢Â§±Ë¥•Êó∂ÈªòËÆ§‰∏∫0.0
            elif output_type == "string":
                converted_result.append(item)
            elif output_type == "boolean":
                # Â§ÑÁêÜÂ∏ÉÂ∞îÂÄºËΩ¨Êç¢
                if item.lower() in ["true", "yes", "1", "on"]:
                    converted_result.append(True)
                elif item.lower() in ["false", "no", "0", "off"]:
                    converted_result.append(False)
                else:
                    converted_result.append(bool(item))  # ÂÖ∂‰ªñÊÉÖÂÜµÊåâÈùûÁ©∫Â≠óÁ¨¶‰∏≤Â§ÑÁêÜ
            elif output_type == "anytype":
                # Â∞ùËØïËá™Âä®ËΩ¨Êç¢Á±ªÂûã
                try:
                    num = int(item)
                    converted_result.append(num)
                except ValueError:
                    try:
                        num = float(item)
                        converted_result.append(num)
                    except ValueError:
                        # Ê£ÄÊü•ÊòØÂê¶‰∏∫Â∏ÉÂ∞îÂÄº
                        if item.lower() in ["true", "yes", "1", "on"]:
                            converted_result.append(True)
                        elif item.lower() in ["false", "no", "0", "off"]:
                            converted_result.append(False)
                        else:
                            converted_result.append(item)
        
        # data_list ËæìÂá∫‰∏é string Á±ªÂûãÁõ∏ÂêåÁöÑÂéüÂßãÂ≠óÁ¨¶‰∏≤ÂàóË°®
        return (converted_result, items)




class IO_load_anyimage:
    @classmethod
    def INPUT_TYPES(s):
        return {
            "required": {
                "file_path": ("STRING", {}),
                "fill_color": (["None", "white", "gray", "black"], {}),
            },
            "optional": {
                "max_images": ("INT", {"default": 0, "min": 0, "max": 1000, "step": 1, 
                                     "tooltip": "0Ë°®Á§∫Êó†ÈôêÂà∂"}),
                "keyword_filter": ("STRING", {"default": "", "multiline": False,
                                            "tooltip": "Âè™ÊúâÊñá‰ª∂ÂêçÂåÖÂê´Ê≠§ÂÖ≥ÈîÆÂ≠óÁöÑÂõæÁâáÊâç‰ºöË¢´Âä†ËΩΩÔºåÁïôÁ©∫Ë°®Á§∫‰∏çËøáÊª§"}),
                "number_prefix": ("BOOLEAN", {"default": False, "label_on": "ÂâçÁΩÆÁºñÂè∑", "label_off": "ÂêéÁΩÆÁºñÂè∑",
                                            "tooltip": "ÂºÄÂêØÊó∂ÊåâÂâçÁΩÆÁºñÂè∑ÊéíÂ∫èÔºåÂÖ≥Èó≠Êó∂ÊåâÂêéÁΩÆÁºñÂè∑ÊéíÂ∫è"}),
                "number_digits": ("INT", {"default": 3, "min": 1, "max": 10, "step": 1,
                                        "tooltip": "ÁºñÂè∑‰ΩçÊï∞ÔºåÂ¶ÇËÆæÁΩÆ‰∏∫3ÂàôËØÜÂà´001Ê†ºÂºèÁöÑÁºñÂè∑"})
            }
        }
    RETURN_TYPES = ('IMAGE', 'MASK',)
    FUNCTION = "get_transparent_image"
    CATEGORY = "Apt_Preset/View_IO"
    
    def extract_number_from_filename(self, filename, number_prefix, number_digits):
        """
        ‰ªéÊñá‰ª∂Âêç‰∏≠ÊèêÂèñÁºñÂè∑
        :param filename: Êñá‰ª∂ÂêçÔºà‰∏çÂê´Êâ©Â±ïÂêçÔºâ
        :param number_prefix: ÊòØÂê¶ÂâçÁΩÆÁºñÂè∑
        :param number_digits: ÁºñÂè∑‰ΩçÊï∞
        :return: ÊèêÂèñÂà∞ÁöÑÁºñÂè∑ÔºåÂ¶ÇÊûúÊú™ÊâæÂà∞ÂàôËøîÂõûNone
        """
        # ÁßªÈô§Êâ©Â±ïÂêç
        name_without_ext = os.path.splitext(filename)[0]
        
        if number_prefix:
            # ÂâçÁΩÆÁºñÂè∑Ê®°ÂºèÔºöÊü•ÊâæÊñá‰ª∂ÂêçÂºÄÂ§¥ÁöÑÊï∞Â≠ó
            pattern = r'^(\d{' + str(number_digits) + r'})'
        else:
            # ÂêéÁΩÆÁºñÂè∑Ê®°ÂºèÔºöÊü•ÊâæÊñá‰ª∂ÂêçÊú´Â∞æÁöÑÊï∞Â≠ó
            pattern = r'(\d{' + str(number_digits) + r'})$'
        
        match = re.search(pattern, name_without_ext)
        if match:
            try:
                return int(match.group(1))
            except ValueError:
                return None
        return None
    
    def get_transparent_image(self, file_path, fill_color, max_images=0, keyword_filter="", 
                             number_prefix=False, number_digits=3):
        try:
            if os.path.isdir(file_path):
                images = []
                image_files_with_numbers = []
                
                # Ëé∑ÂèñÁõÆÂΩï‰∏≠ÊâÄÊúâÁ¨¶ÂêàÊù°‰ª∂ÁöÑÂõæÁâáÊñá‰ª∂
                image_files = [f for f in os.listdir(file_path) 
                              if f.lower().endswith(('.png', '.jpg', '.jpeg', '.webp'))]
                
                # Â¶ÇÊûúËÆæÁΩÆ‰∫ÜÂÖ≥ÈîÆÂ≠óËøáÊª§ÔºåÂàôÂè™‰øùÁïôÂåÖÂê´ÂÖ≥ÈîÆÂ≠óÁöÑÊñá‰ª∂
                if keyword_filter:
                    image_files = [f for f in image_files if keyword_filter in f]
                
                # ÊèêÂèñÊñá‰ª∂ÁºñÂè∑Âπ∂Â≠òÂÇ®Êñá‰ª∂ÂêçÂíåÁºñÂè∑ÁöÑÂÖÉÁªÑ
                for filename in image_files:
                    number = self.extract_number_from_filename(filename, number_prefix, number_digits)
                    if number is not None:
                        image_files_with_numbers.append((filename, number))
                    else:
                        # Â¶ÇÊûúÊ≤°ÊúâÊâæÂà∞ÁºñÂè∑Ôºå‰ΩøÁî®Êñá‰ª∂Âêç‰Ωú‰∏∫ÊéíÂ∫è‰æùÊçÆ
                        image_files_with_numbers.append((filename, filename))
                
                # Ê†πÊçÆÁºñÂè∑ÊéíÂ∫èÔºàÊï∞Â≠ó‰ºòÂÖàÔºåÁÑ∂ÂêéÊòØÂ≠óÁ¨¶‰∏≤Ôºâ
                image_files_with_numbers.sort(key=lambda x: (
                    isinstance(x[1], str),  # Â≠óÁ¨¶‰∏≤ÁºñÂè∑ÊéíÂú®Êï∞Â≠óÁºñÂè∑ÂêéÈù¢
                    x[1]  # ÊåâÁºñÂè∑ÊéíÂ∫è
                ))
                
                # ÊèêÂèñÊéíÂ∫èÂêéÁöÑÊñá‰ª∂ÂêçÂàóË°®
                sorted_image_files = [item[0] for item in image_files_with_numbers]
                
                # Â¶ÇÊûúËÆæÁΩÆ‰∫ÜÊúÄÂ§ßÂõæÁâáÊï∞ÈáèÈôêÂà∂ÔºåÂàôÊà™ÂèñÂâçmax_images‰∏™Êñá‰ª∂
                if max_images > 0:
                    sorted_image_files = sorted_image_files[:max_images]
                
                # Âä†ËΩΩÁ¨¶ÂêàÊù°‰ª∂ÁöÑÂõæÁâá
                for filename in sorted_image_files:
                    img_path = os.path.join(file_path, filename)
                    image = Image.open(img_path).convert('RGBA')
                    images.append(image)
                
                if not images:
                    return None, None
                
                target_size = images[0].size
                
                resized_images = []
                for image in images:
                    if image.size != target_size:
                        image = image.resize(target_size, Image.BILINEAR)
                    resized_images.append(image)
                
                batch_images = np.stack([np.array(img) for img in resized_images], axis=0).astype(np.float32) / 255.0
                batch_tensor = torch.from_numpy(batch_images)
                
                mask_tensor = None
                
                return batch_tensor, mask_tensor        
            else:
                file_path = file_path.strip('"')
                image = Image.open(file_path)
                if image is not None:
                    image_rgba = image.convert('RGBA')
                    # Ê£ÄÊü•Âçï‰∏™Êñá‰ª∂ÊòØÂê¶Á¨¶ÂêàÂÖ≥ÈîÆÂ≠óËøáÊª§Êù°‰ª∂
                    if keyword_filter and keyword_filter not in os.path.basename(file_path):
                        print(f"Êñá‰ª∂ {file_path} ‰∏çÂåÖÂê´ÂÖ≥ÈîÆÂ≠ó '{keyword_filter}'ÔºåË∑≥ËøáÂä†ËΩΩ")
                        return None, None
                    
                    image_rgba.save(file_path.rsplit('.', 1)[0] + '.png')
                       
                    if fill_color == 'white':
                        for y in range(image_rgba.height):
                            for x in range(image_rgba.width):
                                if image_rgba.getpixel((x, y))[3] == 0:
                                    image_rgba.putpixel((x, y), (255, 255, 255, 255))
                    elif fill_color == 'gray':
                        for y in range(image_rgba.height):
                            for x in range(image_rgba.width):
                                if image_rgba.getpixel((x, y))[3] == 0:
                                    image_rgba.putpixel((x, y), (128, 128, 128))
                    elif fill_color == 'black':
                        for y in range(image_rgba.height):
                            for x in range(image_rgba.width):
                                if image_rgba.getpixel((x, y))[3] == 0:
                                    image_rgba.putpixel((x, y), (0, 0, 0))
                    elif fill_color == 'None':
                        pass
                    else:
                        raise ValueError("Invalid fill color specified.")
            
                    image_np = np.array(image_rgba).astype(np.float32) / 255.0
                    image_tensor = torch.from_numpy(image_np)[None, :, :, :]
            
                    return (image_tensor, mask_tensor)
            
        except Exception as e:
            print(f"Âá∫ÈîôËØ∑ÈáçÁΩÆËäÇÁÇπÔºö{e}")
        return None, None








class IO_image_select:
    @classmethod
    def INPUT_TYPES(s):
        return {
            "required": {
                "images": ("IMAGE",),
                "indexes": ("STRING", {"default": "1,2"}),
                "canvas_operations": (["None", "Horizontal Flip", "Vertical Flip", "90 Degree Rotation", 
                                       "180 Degree Rotation", "Horizontal Flip + 90 Degree Rotation", "Horizontal Flip + 180 Degree Rotation"], {"default": "None"}),
            },
            "optional": {
                "index": ("INT", {"forceInput": True}),
            }
        }
    
    CATEGORY = "Apt_Preset/View_IO"
    RETURN_TYPES = ("IMAGE", "IMAGE",)
    RETURN_NAMES = ("select_img", "exclude_img")
    FUNCTION = "SelectImages"
    DESCRIPTION = """indexesÊåâÂõæÂÉèÁ¥¢ÂºïÈÄâÊã©ËæìÂá∫
    Á¥¢ÂºïÊñπÂºèÔºöÊ≠£ ‚Äú1,3,5‚Äù ÈÄÜÂêë ‚Äú-1,-3‚ÄùÔºà-1‰∏∫ÊúÄÂêé‰∏ÄÂº†Ôºâ
    ËåÉÂõ¥ÊñπÂºèÔºö ‚Äú2-4‚Äù  
    Ëã•Êèê‰æõindexÔºàÊï¥Êï∞ÔºâÔºåÂàô‰ºòÂÖàÊåâindexÈÄâÊã©ÂõæÂÉè"""
    
    def parse_indexes(self, indexes_str, max_length):
        indexes = []
        parts = [p.strip() for p in indexes_str.split(',') if p.strip()]
        for part in parts:
            if '-' in part and not part.startswith('-'):
                try:
                    start, end = part.split('-', 1)
                    start = int(start.strip())
                    end = int(end.strip())
                    start = start if start > 0 else max_length + start + 1
                    end = end if end > 0 else max_length + end + 1
                    if start > end:
                        start, end = end, start
                    start = max(1, min(start, max_length))
                    end = max(1, min(end, max_length))
                    indexes.extend(range(start, end + 1))
                except ValueError:
                    continue
            else:
                try:
                    num = int(part)
                    if num < 0:
                        num = max_length + num + 1
                    if 1 <= num <= max_length:
                        indexes.append(num)
                except ValueError:
                    continue
        seen = set()
        unique_indexes = []
        for num in indexes:
            if num not in seen:
                seen.add(num)
                unique_indexes.append(num)
        return unique_indexes
    
    def SelectImages(self, images, indexes, canvas_operations, index=None):
        max_length = len(images)
        if max_length == 0:
            return (images, [])
        
        # ‰ºòÂÖàÂ§ÑÁêÜindexËæìÂÖ•
        if index is not None:
            # ËΩ¨Êç¢‰∏∫1-basedÁ¥¢ÂºïÂπ∂Ê†°È™åËåÉÂõ¥
            if index < 0:
                adjusted_index = max_length + index + 1
            else:
                adjusted_index = index  # indexÊú¨Ë∫´‰Ωú‰∏∫1-basedÁ¥¢ÂºïÂ§ÑÁêÜ
            
            if 1 <= adjusted_index <= max_length:
                select_numbers = [adjusted_index]
            else:
                print(f"Warning: index {index} is out of range, using original indexes logic instead.")
                select_numbers = self.parse_indexes(indexes, max_length)
        else:
            # Ê≤°ÊúâindexËæìÂÖ•Êó∂‰ΩøÁî®ÂéüindexesÈÄªËæë
            select_numbers = self.parse_indexes(indexes, max_length)
        
        if not select_numbers:
            print("Warning: No valid indexes found, return original input.")
            return (images, [])
        
        select_list1 = np.array(select_numbers) - 1  # ËΩ¨Êç¢‰∏∫0-basedÁ¥¢Âºï
        exclude_list = np.setdiff1d(np.arange(max_length), select_list1)
        print(f"Selected images (1-based): {select_numbers}")
        
        selected_images = images[torch.tensor(select_list1, dtype=torch.long)]
        excluded_images = images[torch.tensor(exclude_list, dtype=torch.long)] if len(exclude_list) > 0 else []
        
        def apply_operation(images_tensor):
            if not isinstance(images_tensor, torch.Tensor) or len(images_tensor) == 0:
                return images_tensor
            if canvas_operations == "Horizontal Flip":
                return torch.flip(images_tensor, [2])
            elif canvas_operations == "Vertical Flip":
                return torch.flip(images_tensor, [1])
            elif canvas_operations == "90 Degree Rotation":
                return torch.rot90(images_tensor, 1, [1, 2])
            elif canvas_operations == "180 Degree Rotation":
                return torch.rot90(images_tensor, 2, [1, 2])
            elif canvas_operations == "Horizontal Flip + 90 Degree Rotation":
                flipped = torch.flip(images_tensor, [2])
                return torch.rot90(flipped, 1, [1, 2])
            elif canvas_operations == "Horizontal Flip + 180 Degree Rotation":
                flipped = torch.flip(images_tensor, [2])
                return torch.rot90(flipped, 2, [1, 2])
            return images_tensor
        
        selected_images = apply_operation(selected_images)
        excluded_images = apply_operation(excluded_images)
        
        return (selected_images, excluded_images)





import os
import glob
import json
import toml 
from typing import Tuple


try:
    import docx
    REMOVER_AVAILABLE = True
except ImportError:
    docx = None
    REMOVER_AVAILABLE = False



class XXIO_inputfile:
    
    @classmethod
    def INPUT_TYPES(cls):
        file_types = ["text", "md", "json", "js", "py", "toml"]
        if REMOVER_AVAILABLE:
            file_types.append("docx")
        
        return {
            "required": {
                "path": ("STRING", {
                    "default": "",
                    "placeholder": "ËæìÂÖ•Êñá‰ª∂Ë∑ØÂæÑ"
                }),
                "file_type": (file_types,),
                "char_limit": ("INT", {
                    "default": 0,
                    "min": 0,
                    "max": 100000,
                    "step": 10,
                }),
                "batch_mode": ("BOOLEAN", {
                    "default": False,
                    "label_on": "Êñá‰ª∂Â§πÊâπÈáèËØªÂèñ",
                    "label_off": "ÂçïÊñá‰ª∂ËØªÂèñ"
                }),
            }
        }
    
    RETURN_TYPES = ("STRING", "STRING")
    RETURN_NAMES = ("text", "file_paths")
    FUNCTION = "read_content"
    CATEGORY = "Apt_Preset/View_IO"

    def read_content(self, path: str, file_type: str, char_limit: int, batch_mode: bool) -> Tuple[str, str]:
        path = path.strip('\'"')
        
        if not path:
            raise ValueError("Ë∑ØÂæÑ‰∏çËÉΩ‰∏∫Á©∫")
        
        ext_mapping = {
            "text": ".txt",
            "md": ".md",
            "json": ".json",
            "js": ".js",
            "py": ".py",
            "toml": ".toml",
            "docx": ".docx"
        }
        target_ext = ext_mapping.get(file_type, "")
        
        if file_type == "docx" and not REMOVER_AVAILABLE:
            raise ValueError("Áº∫Â∞ëpython-docxÂ∫ìÔºåËØ∑ÂÆâË£ÖÂêéÈáçËØïÔºàÂèØ‰ΩøÁî®ÂëΩ‰ª§Ôºöpip install python-docxÔºâ")
        
        try:
            if batch_mode:
                if not os.path.isdir(path):
                    raise ValueError(f"ÊâπÈáèÊ®°Âºè‰∏ãË∑ØÂæÑÂøÖÈ°ªÊòØÊñá‰ª∂Â§π - {path}")
                
                search_pattern = os.path.join(path, f"*{target_ext}")
                file_paths = glob.glob(search_pattern)
                
                if not file_paths:
                    raise ValueError(f"Ë≠¶ÂëäÔºöÂú® {path} ‰∏≠Êú™ÊâæÂà∞{target_ext}Á±ªÂûãÊñá‰ª∂")
                
                file_paths.sort(key=lambda x: os.path.basename(x))
                read_paths = []
                
                merged_content = []
                total_char_count = 0
                for file_path in file_paths:
                    file_name = os.path.basename(file_path)
                    merged_content.append(f"\n\n===== ÂºÄÂßãÔºö{file_name} =====")
                    
                    content = self._read_single_file(file_path, file_type)
                    merged_content.append(content)
                    merged_content.append(f"===== ÁªìÊùüÔºö{file_name} =====")
                    
                    read_paths.append(file_path)
                    total_char_count += len(content)
                    
                    if char_limit > 0 and total_char_count > char_limit:
                        merged_content.append(f"\n\n...ÔºàÂ∑≤ËææÂ≠óÁ¨¶ÈôêÂà∂ {char_limit}ÔºåÂêéÁª≠Êñá‰ª∂Êú™ËØªÂèñÔºâ")
                        break
                
                final_content = ''.join(merged_content)
                paths_str = "\n".join(read_paths)
                return (final_content, paths_str)
            
            else:
                if not os.path.isfile(path):
                    raise ValueError(f"Êñá‰ª∂‰∏çÂ≠òÂú® - {path}")
                
                if not path.lower().endswith(target_ext):
                    raise ValueError(f"Ë≠¶ÂëäÔºöÊñá‰ª∂Êâ©Â±ïÂêç‰∏éÊâÄÈÄâÁ±ªÂûã‰∏çÂåπÈÖçÔºàÈ¢ÑÊúü{target_ext}Ôºâ")
                
                content = self._read_single_file(path, file_type)
                
                if char_limit > 0 and len(content) > char_limit:
                    content = content[:char_limit] + f"\n\n...ÔºàÂÜÖÂÆπÂ∑≤Êà™Êñ≠ÔºåÂéüÈïøÂ∫¶{len(content)}Â≠óÁ¨¶Ôºâ"
                
                return (content, path)
                
        except Exception as e:
            raise ValueError(f"ËØªÂèñÂ§±Ë¥•Ôºö{str(e)}")
    
    def _read_single_file(self, file_path: str, file_type: str) -> str:
        if file_type == "json":
            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
                return json.dumps(json_data, ensure_ascii=False, indent=2)
        elif file_type == "toml":
            with open(file_path, 'r', encoding='utf-8') as f:
                toml_data = toml.load(f)
                return toml.dumps(toml_data)
        elif file_type == "docx":
            if not REMOVER_AVAILABLE:
                raise ValueError("python-docxÂ∫ìÊú™ÂÆâË£ÖÔºåÊó†Ê≥ïËØªÂèñdocxÊñá‰ª∂")
            
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()



class IO_outputfile:
    
    @classmethod
    def INPUT_TYPES(cls):
        file_types = ["text", "md", "json", "js", "py", "toml"]
        if REMOVER_AVAILABLE:
            file_types.append("docx")
        
        return {
            "required": {
                "content": ("STRING", {
                    "default": "",
                    "multiline": True,
                    "placeholder": "ËæìÂÖ•Ë¶Å‰øùÂ≠òÁöÑÂÜÖÂÆπ"
                }),
                "file_path": ("STRING", {
                    "default": "",
                    "placeholder": "ËæìÂÖ•Êñá‰ª∂‰øùÂ≠òË∑ØÂæÑÔºàÂåÖÂê´Êñá‰ª∂ÂêçÔºâ"
                }),
                "file_type": (file_types,),
            },
            "optional": {
                "custom_file_name": ("STRING", {
                    "default": "",
                    "placeholder": "Ëá™ÂÆö‰πâÊñá‰ª∂ÂêçÔºà‰∏çÂê´Êâ©Â±ïÂêçÔºåÁïôÁ©∫ÂàôËá™Âä®ÁîüÊàêÔºâ"
                }),
            }
        }
    
    RETURN_TYPES = ("STRING",)
    RETURN_NAMES = ("status",)
    FUNCTION = "write_content"
    CATEGORY = "Apt_Preset/View_IO"

    def write_content(self, content: str, file_path: str, file_type: str, custom_file_name: str = "") -> Tuple[str]:
        file_path = file_path.strip('\'"')
        
        if not file_path:
            raise ValueError("Êñá‰ª∂Ë∑ØÂæÑ‰∏çËÉΩ‰∏∫Á©∫")
        
        ext_mapping = {
            "text": ".txt",
            "md": ".md",
            "json": ".json",
            "js": ".js",
            "py": ".py",
            "toml": ".toml",
            "docx": ".docx"
        }
        target_ext = ext_mapping.get(file_type, "")
        if not target_ext:
            raise ValueError("‰∏çÊîØÊåÅÁöÑÊñá‰ª∂Á±ªÂûã")
        
        try:
            # Â§ÑÁêÜË∑ØÂæÑÂíåÊñá‰ª∂ÂêçÈÄªËæë
            parent_dir = os.path.dirname(file_path)
            original_file_name = os.path.basename(file_path)
            
            # ÊÉÖÂÜµ1ÔºöÊåáÂÆö‰∫ÜËá™ÂÆö‰πâÊñá‰ª∂Âêç
            if custom_file_name.strip():
                # ‰øùÁïôÂéüË∑ØÂæÑÔºå‰ªÖÊõøÊç¢Êñá‰ª∂Âêç
                full_path = os.path.join(parent_dir, f"{custom_file_name.strip()}{target_ext}")
            # ÊÉÖÂÜµ2ÔºöÊú™ÊåáÂÆöËá™ÂÆö‰πâÊñá‰ª∂ÂêçÔºå‰∏îÁõÆÊ†áË∑ØÂæÑ‰∏çÂ≠òÂú®
            elif not os.path.exists(file_path):
                # ÁîüÊàêÈªòËÆ§Êñá‰ª∂ÂêçÔºàÊó∂Èó¥Êà≥ÂâçÁºÄÔºâ
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                default_file_name = f"output_{timestamp}"
                full_path = os.path.join(parent_dir, f"{default_file_name}{target_ext}")
            # ÊÉÖÂÜµ3ÔºöÊú™ÊåáÂÆöËá™ÂÆö‰πâÊñá‰ª∂ÂêçÔºå‰ΩÜË∑ØÂæÑÂ∑≤Â≠òÂú®Ôºà‰ΩøÁî®ÂéüË∑ØÂæÑÔºâ
            else:
                # Á°Æ‰øùÂéüË∑ØÂæÑÂåÖÂê´Ê≠£Á°ÆÊâ©Â±ïÂêç
                full_path = file_path if file_path.lower().endswith(target_ext) else f"{file_path}{target_ext}"
            
            # ÂàõÂª∫Áà∂ÁõÆÂΩïÔºàÂ¶ÇÊûú‰∏çÂ≠òÂú®Ôºâ
            if parent_dir and not os.path.exists(parent_dir):
                os.makedirs(parent_dir, exist_ok=True)
            
            # Ê†πÊçÆÊñá‰ª∂Á±ªÂûãÂÜôÂÖ•ÂÜÖÂÆπ
            if file_type == "json":
                try:
                    json_data = json.loads(content)
                    with open(full_path, 'w', encoding='utf-8') as f:
                        json.dump(json_data, f, ensure_ascii=False, indent=2)
                except (json.JSONDecodeError, TypeError):
                    with open(full_path, 'w', encoding='utf-8') as f:
                        f.write(content)
            
            elif file_type == "toml":
                try:
                    toml_data = toml.loads(content)
                    with open(full_path, 'w', encoding='utf-8') as f:
                        toml.dump(toml_data, f)
                except toml.TomlDecodeError:
                    with open(full_path, 'w', encoding='utf-8') as f:
                        f.write(content)
            
            elif file_type == "docx":
                if not REMOVER_AVAILABLE:
                    raise ValueError("Áº∫Â∞ëpython-docxÂ∫ìÔºåËØ∑ÂÆâË£ÖÂêéÈáçËØï")
                
                doc = docx.Document()
                for para_text in content.split('\n'):
                    doc.add_paragraph(para_text)
                doc.save(full_path)
            
            else:
                with open(full_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            return (f"ÊàêÂäüÔºöÊñá‰ª∂Â∑≤‰øùÂ≠òËá≥ {full_path}",)
        
        except Exception as e:
            raise ValueError(f"ÂÜôÂÖ•Â§±Ë¥•Ôºö{str(e)}")




class view_bridge_image:
    def __init__(self):
        self.image_id = None
        self.cached_mask = None  

    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {
                "image": ("IMAGE",)
            },
            "optional": {
                "mask": ("MASK",), 
                "output_mask": ("BOOLEAN", {"default": False, "label_off": "Refresh Mask", "label_on": "Store Mask"}),  
                "operation": (["+", "-", "*", "&", "None"], {"default": "+"}),
                "image_update": ("IMAGE_FILE",),  

            }
        }

    CATEGORY = "Apt_Preset/View_IO"
    RETURN_TYPES = ("IMAGE", "MASK")
    RETURN_NAMES = ("image", "mask")
    FUNCTION = "edit"
    OUTPUT_NODE = True

    def edit(self, image, mask=None, operation="None", image_update=None, output_mask=False):
        if self.image_id is None:
            self.image_id = tensor_to_hash(image)
            image_update = None
        else:
            image_id = tensor_to_hash(image)
            if image_id != self.image_id:
                image_update = None
                self.image_id = image_id
                # ÂõæÂÉèIDÂèòÂåñÊó∂ÈáçÁΩÆÁºìÂ≠òÈÅÆÁΩ©
                if not output_mask:
                    self.cached_mask = None

        # ‰ºòÂÖà‰ΩøÁî® image_update ‰∏≠ÁöÑÂõæÂÉè
        if image_update is not None and 'images' in image_update:
            images = image_update['images']
            filename = images[0]['filename']
            subfolder = images[0]['subfolder']
            type = images[0]['type']
            name, base_dir = folder_paths.annotated_filepath(filename)

            if type.endswith("output"):
                base_dir = folder_paths.get_output_directory()
            elif type.endswith("input"):
                base_dir = folder_paths.get_input_directory()
            elif type.endswith("temp"):
                base_dir = folder_paths.get_temp_directory()

            image_path = os.path.join(base_dir, subfolder, name)
            img = node_helpers.pillow(Image.open, image_path)
        else:
            # Âê¶Âàô‰ΩøÁî® preview_image
            if mask is not None:
                try:
                    masked_result = generate_masked_black_image(image, mask)
                    preview_image = masked_result["result"][0]
                except Exception as e:
                    print(f"[Error] Failed to apply mask for preview: {e}")
                    preview_image = image
            else:
                preview_image = image

            image_path, images = create_temp_file(preview_image)
            img = node_helpers.pillow(Image.open, image_path)

        # ‰ªéÂõæÂÉè‰∏≠ÊèêÂèñ mask
        output_masks = []
        w, h = None, None
        excluded_formats = ['MPO']

        for i in ImageSequence.Iterator(img):
            i = node_helpers.pillow(ImageOps.exif_transpose, i)
            if i.mode == 'I':
                i = i.point(lambda i: i * (1 / 255))
            image_pil = i.convert("RGB")

            if len(output_masks) == 0:
                w = image_pil.size[0]
                h = image_pil.size[1]

            if image_pil.size[0] != w or image_pil.size[1] != h:
                continue

            if 'A' in i.getbands():
                mask_np = np.array(i.getchannel('A')).astype(np.float32) / 255.0
                mask_tensor = 1. - torch.from_numpy(mask_np)
            else:
                mask_tensor = torch.zeros((h, w), dtype=torch.float32, device="cpu")

            output_masks.append(mask_tensor.unsqueeze(0))

        if len(output_masks) > 1 and img.format not in excluded_formats:
            output_mask_val = torch.cat(output_masks, dim=0)
        else:
            output_mask_val = output_masks[0] if output_masks else torch.zeros_like(image[0, :, :, 0])

        # Êñ∞Â¢û Mask ËøêÁÆóÈÄªËæë
        mask1 = mask
        mask2 = output_mask_val

        # ËÆ°ÁÆóÂΩìÂâçËøêÁÆóÁªìÊûú
        if mask1 is None or operation == "None":
            current_result = mask2
        else:
            invert_mask1 = False
            invert_mask2 = False

            if invert_mask1:
                mask1 = 1 - mask1
            if invert_mask2:
                mask2 = 1 - mask2

            if mask1.dim() == 2:
                mask1 = mask1.unsqueeze(0)
            if mask2.dim() == 2:
                mask2 = mask2.unsqueeze(0)

            b, h, w = image.shape[0], image.shape[1], image.shape[2]
            if mask1.shape != (b, h, w):
                mask1 = torch.zeros((b, h, w), dtype=mask1.dtype, device=mask1.device)
            if mask2.shape != (b, h, w):
                mask2 = torch.zeros((b, h, w), dtype=mask2.dtype, device=mask2.device)

            algorithm = "torch"  # ÁÆÄÂåñÈÄªËæëÔºåÁõ¥Êé•‰ΩøÁî®torch

            if algorithm == "torch":
                if operation == "-":
                    current_result = torch.clamp(mask1 - mask2, min=0, max=1)
                elif operation == "+":
                    current_result = torch.clamp(mask1 + mask2, min=0, max=1)
                elif operation == "*":
                    current_result = torch.clamp(mask1 * mask2, min=0, max=1)
                elif operation == "&":
                    current_result = (torch.round(mask1).bool() & torch.round(mask2).bool()).float()
                else:
                    current_result = mask2  # ÈªòËÆ§Êìç‰Ωú‰∏∫ mask2

        # Ê†πÊçÆoutput_maskÊéßÂà∂ÊòØÂê¶‰øùÁïôÈÅÆÁΩ©
        if output_mask:
            # Â¶ÇÊûúÊòØÁ¨¨‰∏ÄÊ¨°ÂêØÁî®ÂêØÁî®‰øùÁïôÔºåÁºìÂ≠òÂΩìÂâçÁªìÊûú
            if self.cached_mask is None:
                self.cached_mask = current_result
            # ‰ΩøÁî®ÁºìÂ≠òÁöÑÈÅÆÁΩ©‰Ωú‰∏∫ÁªìÊûú
            final_mask = self.cached_mask
        else:
            # ‰∏ç‰øùÁïôÊó∂Êõ¥Êñ∞ÁºìÂ≠ò‰∏∫ÂΩìÂâçÁªìÊûú
            self.cached_mask = current_result
            final_mask = current_result

        # ËøîÂõûÁªìÊûú
        return {"ui": {"images": images}, "result": (image, final_mask)}

    # ‰ª•‰∏ãÈùôÊÄÅÊñπÊ≥ï‰øùÊåÅ‰∏çÂèò
    @staticmethod
    def subtract_masks(mask1, mask2):
        mask1 = mask1.cpu()
        mask2 = mask2.cpu()
        cv2_mask1 = np.array(mask1) * 255
        cv2_mask2 = np.array(mask2) * 255
        if cv2_mask1.shape == cv2_mask2.shape:
            cv2_mask = cv2.subtract(cv2_mask1, cv2_mask2)
            return torch.clamp(torch.from_numpy(cv2_mask) / 255.0, min=0, max=1)
        else:
            print("Warning: The two masks have different shapes")
            return mask1

    @staticmethod
    def add_masks(mask1, mask2):
        mask1 = mask1.cpu()
        mask2 = mask2.cpu()
        cv2_mask1 = np.array(mask1) * 255
        cv2_mask2 = np.array(mask2) * 255
        if cv2_mask1.shape == cv2_mask2.shape:
            cv2_mask = cv2.add(cv2_mask1, cv2_mask2)
            return torch.clamp(torch.from_numpy(cv2_mask) / 255.0, min=0, max=1)
        else:
            print("Warning: The two masks have different shapes")
            return mask1

    @staticmethod
    def multiply_masks(mask1, mask2):
        mask1 = mask1.cpu()
        mask2 = mask2.cpu()
        cv2_mask1 = np.array(mask1) * 255
        cv2_mask2 = np.array(mask2) * 255
        if cv2_mask1.shape == cv2_mask2.shape:
            cv2_mask = cv2.multiply(cv2_mask1, cv2_mask2)
            return torch.clamp(torch.from_numpy(cv2_mask) / 255.0, min=0, max=1)
        else:
            print("Warning: The two masks have different shapes")
            return mask1

    @staticmethod
    def and_masks(mask1, mask2):
        mask1 = mask1.cpu()
        mask2 = mask2.cpu()
        cv2_mask1 = np.array(mask1) * 255
        cv2_mask2 = np.array(mask2) * 255
        if cv2_mask1.shape == cv2_mask2.shape:
            cv2_mask = cv2.bitwise_and(cv2_mask1, cv2_mask2)
            return torch.from_numpy(cv2_mask)
        else:
            print("Warning: The two masks have different shapes")
            return mask1





class view_bridge_Text:    # web_node/view_Data_text.js

    @classmethod
    def INPUT_TYPES(s):
        return {
            "optional": {
                "text": ("STRING", {"forceInput": True}),
                "display": ("STRING", {"default": "", "multiline": True}),
            },
            "hidden": {
                "unique_id": "UNIQUE_ID",
            },
        }

    RETURN_TYPES = ("STRING",)
    RETURN_NAMES = ("text",)

    OUTPUT_NODE = True

    CATEGORY = "Apt_Preset/View_IO"
    FUNCTION = "process"

    def __init__(self):
        self.last_text = None  # Â≠òÂÇ®‰∏ä‰∏ÄÊ¨°ÁöÑ text ËæìÂÖ•ÔºåÁî®‰∫éÊ£ÄÊµãÊõ¥Êñ∞

    def process(self, text="", display="", unique_id=None):
        # Â§ÑÁêÜÂàóË°®ËæìÂÖ•ÔºöÂèñÁ¨¨‰∏Ä‰∏™ÈùûÁ©∫ÂÄºÔºàÂÖºÂÆπÂàóË°®ËæìÂÖ•Âú∫ÊôØÔºâ
        if isinstance(text, list):
            current_text = text[0] if text else ""
        else:
            current_text = text

        # Ëá™Âä®Âà§Êñ≠ edit_modeÔºöÊ£ÄÊµã text ÊòØÂê¶ÊúâÊõ¥Êñ∞
        # ÈÄªËæëÔºöÈ¶ñÊ¨°ËøêË°å/last_text Êú™ÂàùÂßãÂåñ ‚Üí ÈªòËÆ§‰∏∫ TrueÔºõtext ÂÜÖÂÆπÂèòÂåñ ‚Üí FalseÔºõÊó†ÂèòÂåñ ‚Üí True
        if self.last_text is None:
            use_simple = True  # È¶ñÊ¨°Âä†ËΩΩÔºåÊó†ÂéÜÂè≤Êï∞ÊçÆÔºåÈªòËÆ§ÂêØÁî® display Ê®°Âºè
        else:
            use_simple = (current_text == self.last_text)  # text Êó†Êõ¥Êñ∞Âàô‰∏∫ True

        # Êõ¥Êñ∞ÂéÜÂè≤ËÆ∞ÂΩïÔºà‰øùÂ≠òÂΩìÂâç text ‰æõ‰∏ãÊ¨°ÂØπÊØîÔºâ
        self.last_text = current_text

        # Ê®°ÂºèÈÄªËæëÔºöTrue Áî® display ‰Ωú‰∏∫ËæìÂÖ•Ê∫êÔºåFalse Áî® text ‰Ωú‰∏∫ËæìÂÖ•Ê∫ê
        if use_simple:
            input_text = display
        else:
            input_text = current_text

        displayText = self.render(input_text)
        updateTextWidget(unique_id, "display", displayText)
        
        return {"ui": {"display": displayText}, "result": (input_text,)}

    def render(self, input):
        if not isinstance(input, list):
            return input

        listLen = len(input)

        if listLen == 0:
            return ""

        if listLen == 1:
            return input[0]

        result = "List:\n"

        for i, element in enumerate(input):
            result += f"- {element}\n"  # ‰ºòÂåñÔºöÁõ¥Êé•Áî® element ËÄåÈùû input[i]ÔºåÊõ¥ÁÆÄÊ¥Å
        return result





class IO_inputfile:
    
    @classmethod
    def INPUT_TYPES(cls):
        file_types = ["text", "md", "json", "js", "py", "toml"]
        if REMOVER_AVAILABLE:
            file_types.append("docx")
        
        return {
            "required": {
                "path": ("STRING", {
                    "default": "",
                    "placeholder": "ËæìÂÖ•Êñá‰ª∂Ë∑ØÂæÑ"
                }),
                "file_type": (file_types,),
                "char_limit": ("INT", {
                    "default": 0,
                    "min": 0,
                    "max": 100000,
                    "step": 10,
                }),
                "batch_mode": ("BOOLEAN", {
                    "default": False,
                    "label_on": "Êñá‰ª∂Â§πÊâπÈáèËØªÂèñ",
                    "label_off": "ÂçïÊñá‰ª∂ËØªÂèñ"
                }),
            }
        }
    
    RETURN_TYPES = ("STRING", "STRING", "STRING")  # Êñ∞Â¢ûÊñá‰ª∂ÂêçÂ≠óÊÆµËæìÂá∫
    RETURN_NAMES = ("text", "file_paths", "file_names")  # Êñ∞Â¢ûÊñá‰ª∂ÂêçÂ≠óÊÆµÂêçÁß∞
    FUNCTION = "read_content"
    CATEGORY = "Apt_Preset/View_IO"

    def read_content(self, path: str, file_type: str, char_limit: int, batch_mode: bool) -> Tuple[str, str, str]:
        path = path.strip('\'"')
        
        if not path:
            raise ValueError("Ë∑ØÂæÑ‰∏çËÉΩ‰∏∫Á©∫")
        
        ext_mapping = {
            "text": ".txt",
            "md": ".md",
            "json": ".json",
            "js": ".js",
            "py": ".py",
            "toml": ".toml",
            "docx": ".docx"
        }
        target_ext = ext_mapping.get(file_type, "")
        
        if file_type == "docx" and not REMOVER_AVAILABLE:
            raise ValueError("Áº∫Â∞ëpython-docxÂ∫ìÔºåËØ∑ÂÆâË£ÖÂêéÈáçËØïÔºàÂèØ‰ΩøÁî®ÂëΩ‰ª§Ôºöpip install python-docxÔºâ")
        
        try:
            if batch_mode:
                if not os.path.isdir(path):
                    raise ValueError(f"ÊâπÈáèÊ®°Âºè‰∏ãË∑ØÂæÑÂøÖÈ°ªÊòØÊñá‰ª∂Â§π - {path}")
                
                search_pattern = os.path.join(path, f"*{target_ext}")
                file_paths = glob.glob(search_pattern)
                
                if not file_paths:
                    raise ValueError(f"Ë≠¶ÂëäÔºöÂú® {path} ‰∏≠Êú™ÊâæÂà∞{target_ext}Á±ªÂûãÊñá‰ª∂")
                
                file_paths.sort(key=lambda x: os.path.basename(x))
                read_paths = []
                read_names = []  # Â≠òÂÇ®ËØªÂèñÁöÑÊñá‰ª∂ÂêçÂàóË°®
                
                merged_content = []
                total_char_count = 0
                for file_path in file_paths:
                    file_name = os.path.basename(file_path)
                    merged_content.append(f"\n\n===== ÂºÄÂßãÔºö{file_name} =====")
                    
                    content = self._read_single_file(file_path, file_type)
                    merged_content.append(content)
                    merged_content.append(f"===== ÁªìÊùüÔºö{file_name} =====")
                    
                    read_paths.append(file_path)
                    read_names.append(file_name)  # Êî∂ÈõÜÊñá‰ª∂Âêç
                    total_char_count += len(content)
                    
                    if char_limit > 0 and total_char_count > char_limit:
                        merged_content.append(f"\n\n...ÔºàÂ∑≤ËææÂ≠óÁ¨¶ÈôêÂà∂ {char_limit}ÔºåÂêéÁª≠Êñá‰ª∂Êú™ËØªÂèñÔºâ")
                        break
                
                final_content = ''.join(merged_content)
                paths_str = "\n".join(read_paths)
                names_str = "\n".join(read_names)  # Êñá‰ª∂ÂêçÁî®Êç¢Ë°åÂàÜÈöîÊãºÊé•
                return (final_content, paths_str, names_str)
            
            else:
                if not os.path.isfile(path):
                    raise ValueError(f"Êñá‰ª∂‰∏çÂ≠òÂú® - {path}")
                
                if not path.lower().endswith(target_ext):
                    raise ValueError(f"Ë≠¶ÂëäÔºöÊñá‰ª∂Êâ©Â±ïÂêç‰∏éÊâÄÈÄâÁ±ªÂûã‰∏çÂåπÈÖçÔºàÈ¢ÑÊúü{target_ext}Ôºâ")
                
                content = self._read_single_file(path, file_type)
                file_name = os.path.basename(path)  # Ëé∑ÂèñÂçï‰∏™Êñá‰ª∂ÁöÑÊñá‰ª∂Âêç
                
                if char_limit > 0 and len(content) > char_limit:
                    content = content[:char_limit] + f"\n\n...ÔºàÂÜÖÂÆπÂ∑≤Êà™Êñ≠ÔºåÂéüÈïøÂ∫¶{len(content)}Â≠óÁ¨¶Ôºâ"
                
                return (content, path, file_name)  # ËøîÂõûÂçï‰∏™Êñá‰ª∂Âêç
                
        except Exception as e:
            raise ValueError(f"ËØªÂèñÂ§±Ë¥•Ôºö{str(e)}")
    
    def _read_single_file(self, file_path: str, file_type: str) -> str:
        if file_type == "json":
            with open(file_path, 'r', encoding='utf-8') as f:
                json_data = json.load(f)
                return json.dumps(json_data, ensure_ascii=False, indent=2)
        elif file_type == "toml":
            with open(file_path, 'r', encoding='utf-8') as f:
                toml_data = toml.load(f)
                return toml.dumps(toml_data)
        elif file_type == "docx":
            if not REMOVER_AVAILABLE or docx is None:
                raise ValueError("python-docxÂ∫ìÊú™ÂÆâË£ÖÔºåÊó†Ê≥ïËØªÂèñdocxÊñá‰ª∂")
            
            doc = docx_Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        else:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()





class IO_getFilePath:
    @classmethod
    def INPUT_TYPES(cls) -> dict:
        comfy_folder_options = list(folder_paths.folder_names_and_paths.keys()) if (folder_paths and hasattr(folder_paths, 'folder_names_and_paths')) else []
        return {
            "required": {
                "folder_path": ("STRING", {
                    "multiline": False,
                    "default": "",
                    "placeholder": "ÁªùÂØπË∑ØÂæÑÊàñÁõ∏ÂØπË∑ØÂæÑ"
                }),
                "recursive": ("BOOLEAN", {
                    "default": True,
                    "label_on": "ÈÄíÂΩí",
                    "label_off": "ÈùûÈÄíÂΩí",
                    "tooltip": "ÈÄíÂΩíÔºåÂ±ïÂºÄÊâÄÊúâÊñá‰ª∂Â§πÈáåÁöÑÊñá‰ª∂"
                }),
                "file_extensions": ("STRING", {
                    "multiline": False,
                    "default": "",
                    "tooltip": "ËøáÊª§Êâ©Â±ïÂêçÔºà‰æãÔºö.png,.jpg Â§ö‰∏™Áî®ÈÄóÂè∑ÂàÜÈöîÔºåÁïôÁ©∫ÂåπÈÖçÊâÄÊúâÊñá‰ª∂Ôºâ"
                }),
            },
            "optional": {
                "comfy_folder": (comfy_folder_options,),
            }
        }

    RETURN_TYPES = ("STRING", "LIST",)
    RETURN_NAMES = ("Êñá‰ª∂Ë∑ØÂæÑ", "Ë∑ØÂæÑÂàóË°®",)
    FUNCTION = "get_file_paths"
    CATEGORY = "Apt_Preset/View_IO"
    OUTPUT_IS_LIST = (False, True)

    def get_file_paths(
        self,
        folder_path: str,
        recursive: bool = True,
        file_extensions: str = "",
        comfy_folder: Optional[str] = None
    ) -> Tuple[str, List[str]]:
        folder_path = folder_path.strip().strip('"')
        
        if comfy_folder and folder_paths and hasattr(folder_paths, 'folder_names_and_paths'):
            comfy_folder_data = folder_paths.folder_names_and_paths.get(comfy_folder)
            if comfy_folder_data and isinstance(comfy_folder_data, (list, tuple)) and len(comfy_folder_data) > 0:
                base_paths = comfy_folder_data[0]
                if isinstance(base_paths, (list, tuple)) and len(base_paths) > 0:
                    base_path = base_paths[0]
                else:
                    base_path = base_paths

                if isinstance(base_path, str) and folder_path.strip():
                    folder_path = os.path.join(base_path, folder_path.strip())
                elif isinstance(base_path, str):
                    folder_path = base_path

        if folder_path.strip():
            if not os.path.isabs(folder_path):
                base_dir = folder_paths.base_path if (folder_paths and hasattr(folder_paths, 'base_path')) else ""
                folder_path = os.path.abspath(os.path.join(base_dir, folder_path))

            if not os.path.isdir(folder_path):
                print(f"Ë≠¶ÂëäÔºöÊñá‰ª∂Â§π‰∏çÂ≠òÂú®Êàñ‰∏çÊòØÁõÆÂΩïÔºö{folder_path}")
                return ("", [])
        else:
            print("Ë≠¶ÂëäÔºöÊñá‰ª∂Â§πË∑ØÂæÑ‰∏∫Á©∫")
            return ("", [])

        extensions = []
        if file_extensions.strip():
            extensions = [ext.strip().lower() for ext in file_extensions.split(",") if ext.strip()]

        file_paths = []

        try:
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    if extensions:
                        file_ext = os.path.splitext(file)[1].lower()
                        if file_ext not in extensions:
                            continue

                    full_path = os.path.abspath(os.path.join(root, file))
                    file_paths.append(full_path)

                if not recursive:
                    break
        except Exception as e:
            print(f"Ëé∑ÂèñÊñá‰ª∂Ë∑ØÂæÑÂ§±Ë¥•Ôºö{str(e)}")
            return ("", [])

        file_paths.sort()
        paths_string = "\n".join(file_paths) if file_paths else ""
        paths_list = file_paths

        return (paths_string, paths_list)












#region-------------IO_store_image-------------

import torch
import numpy as np
import io
import base64
import json
from PIL import Image
from typing import Optional, Dict, Any, List

# ÂÖ®Â±ÄÂ≠òÂÇ®ÂèòÈáèÔºàËäÇÁÇπÈáçÂª∫Êó∂‰ºöË¢´ __init__ Ê∏ÖÁ©∫Ôºâ
GLOBAL_STORED_IMAGES: List[torch.Tensor] = []
GLOBAL_DISPLAY_DATA: List[Dict[str, Any]] = []

class IO_store_image:
    @classmethod
    def INPUT_TYPES(cls):
        return {
            "required": {},
            "optional": {
                "image": ("IMAGE", {}),
                "output_mode": ("BOOLEAN", {"default": False, "label_on": "last_image", "label_off": "all_image"}),
                "image_output": (["Hide", "Preview", "Save", "Hide/Save"], {"default": "Preview"}),
                "release_total": ("INT", {"default": 0, "min": 0, "step": 1}),
            },
            "hidden": {"prompt": "PROMPT", "extra_pnginfo": "EXTRA_PNGINFO",},
        }
    
    RETURN_TYPES = ("IMAGE", "INT")
    RETURN_NAMES = ("image", "total")
    FUNCTION = "store_image"
    CATEGORY = "Apt_Preset/View_IO"
    OUTPUT_NODE = True
    DESCRIPTION = """
    release_storageÔºàÊ∏ÖÁ©∫ÊâÄÊúâÂ≠òÂÇ®ÔºâÔºöÊ∏ÖÁ©∫ÊâÄÊúâÂ≠òÂÇ®ÁöÑÂõæÂÉèÔºåËæìÂÖ•ÂõæÁâáÁõ¥Êé•ËæìÂá∫‰∏çÂ≠òÂÇ®
    storage_allÔºà‰∏¥Êó∂Â≠òÂÇ®ÔºâÔºö‰ªÖÂú®ÂΩìÂâçËøêË°å‰∏≠Â≠òÂÇ®ÂõæÁâáÔºåÈáçÊñ∞Âä†ËΩΩËäÇÁÇπÊó∂Ê∏ÖÁ©∫ÊâÄÊúâÂ≠òÂÇ®
    ËæìÂá∫ÈÄªËæëÔºö
    - ÂΩì release_total Êú™ËæìÂÖ•ÔºàÊàñ‚â§0ÔºâÊó∂ÔºöËæìÂá∫ÂÖ®ÈÉ®Â≠òÂÇ®ÁöÑÂõæÁâáÔºàÂéüÈÄªËæëÔºâ
    - ÂΩì release_total Â∑≤ËæìÂÖ•Ôºà>0ÔºâÊó∂Ôºö
      1. totalÔºàÂ≠òÂÇ®ÊÄªÊï∞Ôºâ= release_total ‚Üí ËæìÂá∫ÂÖ®ÈÉ®Â≠òÂÇ®ÁöÑÂõæÁâá
      2. total < release_total ‚Üí ËæìÂá∫ÊÄªÊï∞Áî®ÁôΩÂ∫ïÂõæË°•ÈΩêÊï∞Èáè
      3. total > release_total ‚Üí ËæìÂá∫ÂÄíÊï∞ release_total Âº†Â≠òÂÇ®ÁöÑÂõæÁâá
    Ê≥®ÊÑèÔºöÂ≠òÂÇ®Ê®°Âºè‰∏ãÔºå‰ªÖÂΩìËæìÂÖ•‰∏éÊúÄÂêé‰∏ÄÂº†Â≠òÂÇ®ÂõæÂÉè‰∏çÂêåÊó∂ÊâçÊ∑ªÂä†"""

    def __init__(self):
        # ËäÇÁÇπÂàõÂª∫Êó∂Âº∫Âà∂Ê∏ÖÁ©∫ÂÖ®Â±ÄÂ≠òÂÇ®ÔºàÊ†∏ÂøÉÔºö‰øÆÂ§çÊåâÈíÆÈáçÂª∫ËäÇÁÇπÊó∂‰ºöËß¶ÂèëÊ≠§ÊñπÊ≥ïÔºâ
        global GLOBAL_STORED_IMAGES, GLOBAL_DISPLAY_DATA
        GLOBAL_STORED_IMAGES = []
        GLOBAL_DISPLAY_DATA = []
        print("IO_store_image node initialized (storage reset)")

    def store_image(self, image: Optional[torch.Tensor] = None, output_mode: bool = False, 
                   prompt: Any = None, image_output: str = None, 
                   extra_pnginfo: Any = None, release_total: float = 0) -> Dict[str, Any]:
        global GLOBAL_STORED_IMAGES, GLOBAL_DISPLAY_DATA
        
        total = len(GLOBAL_STORED_IMAGES)
        
        if output_mode:
            GLOBAL_STORED_IMAGES = []
            GLOBAL_DISPLAY_DATA = []
            output_image = image
        else:
            if image is not None:
                if not GLOBAL_STORED_IMAGES or not torch.allclose(image, GLOBAL_STORED_IMAGES[-1]):
                    GLOBAL_STORED_IMAGES.append(image)
                    GLOBAL_DISPLAY_DATA.append(self._prepare_image_display(image))
                    total = len(GLOBAL_STORED_IMAGES)
            
            if release_total <= 0:
                output_image = torch.cat(GLOBAL_STORED_IMAGES, dim=0) if GLOBAL_STORED_IMAGES else image
            else:
                release_total = int(release_total)
                if total == 0:
                    output_image = image
                else:
                    if total == release_total:
                        output_image = torch.cat(GLOBAL_STORED_IMAGES, dim=0)
                    elif total < release_total:
                        first_img = GLOBAL_STORED_IMAGES[0]
                        batch_size, channels, height, width = first_img.shape
                        
                        if batch_size > 1:
                            first_img = first_img[0:1]
                        
                        white_img = torch.ones_like(first_img)
                        need_white = release_total - total
                        output_image = torch.cat(
                            GLOBAL_STORED_IMAGES + [white_img] * need_white,
                            dim=0
                        )
                    else:
                        start_idx = total - release_total
                        output_image = torch.cat(GLOBAL_STORED_IMAGES[start_idx:], dim=0)
        
        # ÂÅáËÆæ easySave Â∑≤Âú®ÂÖ∂‰ªñÂú∞ÊñπÂÆö‰πâÔºåËã•Êú™ÂÆö‰πâÈúÄË°•ÂÖÖÂÆûÁé∞
        try:
            results = easySave(output_image, 'easyPreview', image_output, prompt, extra_pnginfo)
        except NameError:
            # Ëã• easySave Êú™ÂÆö‰πâÔºå‰∏¥Êó∂ËøîÂõûÁ©∫ÂàóË°®ÔºàÈÅøÂÖçÊä•ÈîôÔºâ
            results = []
            print("Warning: easySave function not found")
        
        current_total = len(GLOBAL_STORED_IMAGES)
        
        if image_output in ("Hide", "Hide/Save"):
            return {"ui": {}, "result": (output_image, current_total)}
        return {"ui": {"images": results}, "result": (output_image, current_total)}

    def _prepare_image_display(self, image_tensor: torch.Tensor) -> Dict[str, Any]:
        try:
            img_np = image_tensor[0].cpu().numpy()
            img_np = np.transpose(img_np, (1, 2, 0))
            img_np = (img_np * 255).astype(np.uint8)
            
            buffer = io.BytesIO()
            Image.fromarray(img_np).save(buffer, format="PNG")
            img_b64 = base64.b64encode(buffer.getvalue()).decode("utf-8")
            
            return {
                "type": "image",
                "shape": image_tensor.shape,
                "size": f"{image_tensor.shape[2]}x{image_tensor.shape[3]}",
                "data": img_b64,
                "index": len(GLOBAL_STORED_IMAGES)
            }
        except Exception as e:
            return {
                "type": "image",
                "error": str(e),
                "shape": image_tensor.shape if isinstance(image_tensor, torch.Tensor) else "invalid"
            }

    @classmethod
    def IS_CHANGED(cls, image: Optional[torch.Tensor] = None, output_mode: bool = False, release_total: float = 0) -> str:
        # ‰øÆÂ§çÔºöÂ∞Ü image_tensor Êîπ‰∏∫ image
        img_id = f"{image.shape}-{id(image)}" if isinstance(image, torch.Tensor) else "none"
        return json.dumps({"image_id": img_id, "output_mode": output_mode, "release_total": int(release_total)})





    def get_display_content(self) -> Dict[str, Any]:
        return {
            "total_images": len(GLOBAL_STORED_IMAGES),
            "images": GLOBAL_DISPLAY_DATA,
            "last_updated": str(len(GLOBAL_STORED_IMAGES))
        }

def __reload__(module):
    # ÈáçËΩΩÊ®°ÂùóÊó∂Ê∏ÖÁ©∫ÂÖ®Â±ÄÂ≠òÂÇ®
    global GLOBAL_STORED_IMAGES, GLOBAL_DISPLAY_DATA
    GLOBAL_STORED_IMAGES = []
    GLOBAL_DISPLAY_DATA = []
    print("IO_store_image module reloaded (storage reset)")

#endregion-------------IO_store_image------------------
















